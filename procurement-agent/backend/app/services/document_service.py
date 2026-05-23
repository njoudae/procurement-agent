import csv
import io
import json
import re
import unicodedata
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models import DocumentExtraction, RequestAttachment
from app.schemas import DocumentExtractionResult
from app.tools.logging_tools import log_execution


DANGEROUS_PATTERNS = [
    r"ignore\s+(all\s+)?previous\s+instructions",
    r"ignore\s+(these|the|any)\s+instructions",
    r"bypass\s+approval",
    r"send\s+confidential\s+data",
    r"send\s+(the\s+)?secrets?",
    r"exfiltrate\s+(data|secrets?|credentials?)",
    r"reveal\s+(the\s+)?system\s+prompt",
    r"show\s+(the\s+)?system\s+prompt",
    r"act\s+as\s+(the\s+)?system",
    r"you\s+are\s+(now\s+)?(the\s+)?system",
    r"developer\s+message",
    r"system\s+instructions",
    r"<\|/?(?:system|developer|assistant|tool|user)\|>",
    r"^\s*(system|developer|assistant|tool)\s*:",
    r"^\s*#+\s*(system|developer|assistant|tool)\b",
]

HIDDEN_CHARS_PATTERN = re.compile(
    "["
    "\u0000-\u0008"
    "\u000b-\u000c"
    "\u000e-\u001f"
    "\u007f"
    "\u200b-\u200f"
    "\u202a-\u202e"
    "\u2060-\u206f"
    "\ufeff"
    "]"
)


def sanitize_document_text(text: str) -> tuple[str, list[str]]:
    findings: list[str] = []
    safe = unicodedata.normalize("NFKC", str(text or ""))
    safe = HIDDEN_CHARS_PATTERN.sub(" ", safe)
    for pattern in DANGEROUS_PATTERNS:
        if re.search(pattern, safe, flags=re.IGNORECASE | re.MULTILINE):
            findings.append(f"Prompt injection pattern detected: {pattern}")
            safe = re.sub(
                pattern,
                "[REMOVED_UNTRUSTED_INSTRUCTION]",
                safe,
                flags=re.IGNORECASE | re.MULTILINE,
            )
    safe = re.sub(r"\s+", " ", safe).strip()
    return safe[: get_settings().max_document_chars], findings


def sanitize_extracted_tables(tables: list[list[dict[str, Any]]]) -> tuple[list[list[dict[str, Any]]], list[str]]:
    sanitized_tables: list[list[dict[str, Any]]] = []
    findings: list[str] = []
    for table in tables:
        sanitized_rows: list[dict[str, Any]] = []
        for row in table:
            sanitized_row: dict[str, Any] = {}
            for raw_key, raw_value in row.items():
                key, key_findings = sanitize_document_text(str(raw_key))
                value, value_findings = sanitize_document_text(str(raw_value))
                findings.extend(key_findings)
                findings.extend(value_findings)
                sanitized_row[key or "column"] = value
            sanitized_rows.append(sanitized_row)
        sanitized_tables.append(sanitized_rows)
    return sanitized_tables, findings


def extract_attachment(db: Session, attachment: RequestAttachment) -> DocumentExtractionResult:
    try:
        if attachment.SourceType == "pdf":
            result = _extract_pdf(Path(attachment.StoredPath))
        elif attachment.SourceType in {"excel", "csv"}:
            result = _extract_tabular(Path(attachment.StoredPath), attachment.SourceType)
        elif attachment.SourceType == "docx":
            result = _extract_docx(Path(attachment.StoredPath))
        elif attachment.SourceType == "txt":
            result = _extract_txt(Path(attachment.StoredPath))
        elif attachment.SourceType == "image":
            result = DocumentExtractionResult(
                source_type="image",
                source_file=attachment.OriginalFileName,
                extraction_errors=["OCR_REQUIRED: image OCR pipeline placeholder"],
                extraction_confidence=0,
                requires_review=True,
            )
        else:
            result = DocumentExtractionResult(
                source_type=attachment.SourceType,
                source_file=attachment.OriginalFileName,
                extraction_errors=["Unsupported file type"],
                requires_review=True,
            )
    except Exception as exc:
        result = DocumentExtractionResult(
            source_type=attachment.SourceType,
            source_file=attachment.OriginalFileName,
            extraction_errors=[str(exc)],
            requires_review=True,
            extraction_confidence=0,
        )

    result.source_file = attachment.OriginalFileName
    injection_findings: list[str] = []
    result.extracted_text, text_findings = sanitize_document_text(result.extracted_text)
    result.extracted_tables, table_findings = sanitize_extracted_tables(result.extracted_tables)
    injection_findings.extend(text_findings)
    injection_findings.extend(table_findings)
    if injection_findings:
        result.extraction_errors.extend(injection_findings)
        result.requires_review = True
        log_execution(db, attachment.RequestID, "document_guardrail", "NeedsReview", "; ".join(injection_findings))

    extraction = DocumentExtraction(
        RequestID=attachment.RequestID,
        AttachmentID=attachment.AttachmentID,
        SourceType=result.source_type,
        SourceFile=result.source_file,
        ExtractedText=result.extracted_text,
        ExtractedTables=json.dumps(result.extracted_tables, default=str),
        StructuredData=json.dumps(
            {
                "detected_vendor_name": result.detected_vendor_name,
                "detected_quotation_number": result.detected_quotation_number,
                "detected_total_amount": result.detected_total_amount,
                "detected_delivery_date": result.detected_delivery_date,
                "detected_validity_period": result.detected_validity_period,
                "extracted_items": result.extracted_items,
                "extracted_prices": result.extracted_prices,
            },
            default=str,
        ),
        ExtractionConfidence=result.extraction_confidence,
        ExtractionErrors=json.dumps(result.extraction_errors),
        RequiresReview=result.requires_review,
    )
    attachment.ExtractionStatus = "NeedsReview" if result.requires_review else "Extracted"
    db.add(extraction)
    db.commit()
    db.refresh(extraction)
    return result


def extract_pending_attachments(db: Session, request_id: int) -> list[DocumentExtractionResult]:
    attachments = db.scalars(
        select(RequestAttachment).where(RequestAttachment.RequestID == request_id).order_by(RequestAttachment.UploadedAt.asc())
    ).all()
    results: list[DocumentExtractionResult] = []
    seen_hashes: set[str] = set()
    for attachment in attachments:
        if attachment.FileHash in seen_hashes:
            attachment.ExtractionStatus = "Duplicate"
            db.commit()
            log_execution(db, request_id, "extract_attachment_content", "NeedsReview", f"Duplicate attachment {attachment.OriginalFileName}")
            continue
        seen_hashes.add(attachment.FileHash)
        results.append(extract_attachment(db, attachment))
    return results


def _extract_pdf(path: Path) -> DocumentExtractionResult:
    text_parts: list[str] = []
    tables: list[list[dict[str, Any]]] = []
    errors: list[str] = []

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
                for table in page.extract_tables() or []:
                    tables.append(_table_to_dicts(table))
    except Exception as exc:
        errors.append(f"pdfplumber failed: {exc}")

    if not "".join(text_parts).strip():
        try:
            import fitz

            with fitz.open(path) as document:
                if document.needs_pass:
                    raise ValueError("Password-protected PDF")
                text_parts = [page.get_text("text") for page in document]
        except Exception as exc:
            errors.append(f"pymupdf failed: {exc}")

    if not "".join(text_parts).strip():
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            if reader.is_encrypted:
                raise ValueError("Password-protected PDF")
            text_parts = [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:
            errors.append(f"pypdf failed: {exc}")

    text = "\n".join(text_parts).strip()
    if not text:
        errors.append("OCR_REQUIRED: no readable text found in PDF")
    return _build_result("pdf", text, tables, errors)


def _extract_tabular(path: Path, source_type: str) -> DocumentExtractionResult:
    tables: list[list[dict[str, Any]]] = []
    errors: list[str] = []
    try:
        import pandas as pd

        if source_type == "csv":
            frames = {"csv": pd.read_csv(path)}
        else:
            frames = pd.read_excel(path, sheet_name=None)
        for sheet_name, frame in frames.items():
            normalized = frame.rename(columns=lambda col: str(col).strip().lower().replace(" ", "_"))
            rows = normalized.fillna("").to_dict(orient="records")
            tables.append([{"sheet": sheet_name, **row} for row in rows[:100]])
    except Exception as exc:
        errors.append(f"tabular extraction failed: {exc}")
        if source_type == "csv":
            try:
                rows = list(csv.DictReader(io.StringIO(path.read_text(encoding="utf-8", errors="ignore"))))
                tables.append(rows[:100])
            except Exception as csv_exc:
                errors.append(f"csv fallback failed: {csv_exc}")
    text = "\n".join(json.dumps(table[:20], default=str) for table in tables)
    return _build_result(source_type, text, tables, errors)


def _extract_docx(path: Path) -> DocumentExtractionResult:
    errors: list[str] = []
    tables: list[list[dict[str, Any]]] = []
    try:
        from docx import Document

        document = Document(str(path))
        text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(_table_to_dicts(rows))
        return _build_result("docx", "\n".join(text_parts), tables, errors)
    except Exception as exc:
        errors.append(f"docx extraction failed: {exc}")
        return _build_result("docx", "", tables, errors)


def _extract_txt(path: Path) -> DocumentExtractionResult:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return _build_result("txt", text, [], [])
    except Exception as exc:
        return _build_result("txt", "", [], [f"text extraction failed: {exc}"])


def _table_to_dicts(rows: list[list[Any]]) -> list[dict[str, Any]]:
    if not rows:
        return []
    headers = [str(cell or f"column_{idx + 1}").strip().lower().replace(" ", "_") for idx, cell in enumerate(rows[0])]
    output = []
    for row in rows[1:101]:
        output.append({headers[idx] if idx < len(headers) else f"column_{idx + 1}": value for idx, value in enumerate(row)})
    return output


def _build_result(source_type: str, text: str, tables: list[list[dict[str, Any]]], errors: list[str]) -> DocumentExtractionResult:
    safe_text, text_findings = sanitize_document_text(text)
    safe_tables, table_findings = sanitize_extracted_tables(tables)
    safe_errors = list(errors) + text_findings + table_findings
    structured = _detect_quotation_details(safe_text, safe_tables)
    confidence = 0.35
    if safe_text.strip():
        confidence += 0.25
    if safe_tables:
        confidence += 0.25
    if structured["extracted_items"] or structured["extracted_prices"]:
        confidence += 0.1
    confidence = min(confidence, 0.95)
    return DocumentExtractionResult(
        source_type=source_type,
        extracted_text=safe_text,
        extracted_tables=safe_tables,
        extraction_confidence=0 if safe_errors and not safe_text and not safe_tables else confidence,
        extraction_errors=safe_errors,
        requires_review=bool(safe_errors),
        **structured,
    )


def _detect_quotation_details(text: str, tables: list[list[dict[str, Any]]]) -> dict[str, Any]:
    vendor = _regex_first(text, [r"vendor\s*[:#-]\s*([A-Za-z0-9 &.,-]+)", r"supplier\s*[:#-]\s*([A-Za-z0-9 &.,-]+)"])
    quote = _regex_first(text, [r"quotation\s*(?:number|no\.?|#)?\s*[:#-]\s*([A-Za-z0-9-]+)", r"quote\s*(?:number|no\.?|#)?\s*[:#-]\s*([A-Za-z0-9-]+)"])
    delivery_date = _regex_first(text, [r"delivery\s+date\s*[:#-]\s*([A-Za-z0-9 ,./-]+)", r"required\s+date\s*[:#-]\s*([A-Za-z0-9 ,./-]+)"])
    validity = _regex_first(text, [r"validity\s*(?:period)?\s*[:#-]\s*([A-Za-z0-9 ,./-]+)", r"valid\s+until\s*[:#-]\s*([A-Za-z0-9 ,./-]+)"])
    prices = []
    for amount in re.findall(r"(?:USD|SAR|\$)?\s?([0-9][0-9,]*(?:\.[0-9]{2})?)", text):
        prices.append({"value": amount.replace(",", ""), "source": "text"})
    total_amount = None
    total_match = re.search(r"(?:total|grand\s+total|amount\s+due)\s*[:#-]?\s*(?:USD|SAR|\$)?\s?([0-9][0-9,]*(?:\.[0-9]{2})?)", text, flags=re.IGNORECASE)
    if total_match:
        total_amount = float(total_match.group(1).replace(",", ""))
    items = []
    for table in tables:
        for row in table[:50]:
            row_text = " ".join(str(value) for value in row.values())
            if any(key in row for key in ["item", "description", "quantity", "qty", "price", "unit_price", "total"]):
                items.append(row)
            elif re.search(r"\b(qty|quantity|price|total)\b", row_text, flags=re.IGNORECASE):
                items.append(row)
    return {
        "detected_vendor_name": vendor,
        "detected_quotation_number": quote,
        "detected_total_amount": total_amount,
        "detected_delivery_date": delivery_date,
        "detected_validity_period": validity,
        "extracted_items": items[:50],
        "extracted_prices": prices[:50],
    }


def _regex_first(text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()[:255]
    return None
